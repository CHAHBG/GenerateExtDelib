"""
Moteur de génération des Extraits de Délibération - Version Web (Pyodide)
Adapté pour fonctionner dans un navigateur via WebAssembly.
"""

import pandas as pd
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os
import re
import shutil

# === CHEMINS VIRTUELS PYODIDE ===
INPUT_DIR = "/input"
OUTPUT_DIR = "/output"

# Fichiers attendus
FILE_INDIV_DELIB = f"{INPUT_DIR}/INDIV.xlsx"
FILE_COLL_DELIB = f"{INPUT_DIR}/COLL.xlsx"
FILE_COORD_PI = f"{INPUT_DIR}/COORDS_PI.xlsx"
FILE_COORD_PC = f"{INPUT_DIR}/COORDS_PC.xlsx"
TEMPLATE_INDIV = f"{INPUT_DIR}/Template_Indiv.docx"
TEMPLATE_COLL = f"{INPUT_DIR}/Template_Coll.docx"

def log(msg):
    print(msg)

def clean_id(x):
    """Nettoie les identifiants NICAD pour éviter les problèmes de correspondance (123.0 vs 123)"""
    if pd.isnull(x): return ""
    try:
        if isinstance(x, float) and x.is_integer():
            return str(int(x)).strip()
        return str(x).strip()
    except:
        return str(x).strip()

def nettoyer_connexions_donnees(doc):
    """Supprime les connexions de données et paramètres de publipostage"""
    try:
        settings = doc.settings
        element = settings.element
        mail_merge = element.find(qn('w:mailMerge'))
        if mail_merge is not None:
            element.remove(mail_merge)
        web_settings = element.find(qn('w:webSettings'))
        if web_settings is not None:
            element.remove(web_settings)
    except:
        pass

def optimiser_mise_en_page(doc):
    for section in doc.sections:
        section.top_margin = Cm(1.27)
        section.bottom_margin = Cm(1.27)
        section.left_margin = Cm(1.27)
        section.right_margin = Cm(1.27)

def optimiser_espacement(doc):
    for para in doc.paragraphs:
        if para.paragraph_format.space_after and para.paragraph_format.space_after > Pt(6):
            para.paragraph_format.space_after = Pt(3)
        para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

def reduire_texte_legal(doc):
    mots_cles = ['CERTIFIÉ CONFORME', 'APPROUVEE', 'SOUS-PREFET', 'LE MAIRE', 'FAIT LE',
                 'arrêté préfectoral', 'délibération a été approuvée']
    for para in doc.paragraphs:
        if any(mot in para.text for mot in mots_cles):
            for run in para.runs:
                run.font.size = Pt(9)
                run.font.name = 'Times New Roman'

def remplacer_texte(doc, replacements):
    for para in doc.paragraphs:
        texte = para.text
        keys_in_text = [k for k in replacements if k in texte]
        
        if not keys_in_text:
            continue
            
        # Gestion Article 1
        if "Article 1" in texte and para.runs:
            texte_modifie = texte
            for k, v in replacements.items():
                if k in texte_modifie:
                    val = str(v) if pd.notnull(v) else ""
                    texte_modifie = texte_modifie.replace(k, val)
            
            for run in para.runs: run.text = ""
            
            parts = texte_modifie.split(":", 1)
            if len(parts) > 1:
                r1 = para.add_run(parts[0] + ":")
                r1.font.name = 'Times New Roman'
                r1.font.size = Pt(12)
                r1.font.bold = True
                r1.font.underline = True
                
                r2 = para.add_run(parts[1])
                r2.font.name = 'Times New Roman'
                r2.font.size = Pt(11)
                r2.font.bold = True
                r2.font.underline = False
            else:
                para.add_run(texte_modifie)
            continue

        # Gestion générale avec mise en gras
        keys_sorted = sorted(keys_in_text, key=len, reverse=True)
        pattern = '|'.join(map(re.escape, keys_sorted))
        segments = re.split(f'({pattern})', texte)
        
        for run in para.runs:
            run.text = ""
        
        for seg in segments:
            if seg in replacements:
                val = replacements[seg]
                val_str = str(val) if pd.notnull(val) else ""
                run = para.add_run(val_str)
                run.font.bold = True
                run.font.name = 'Times New Roman'
            else:
                run = para.add_run(seg)
                run.font.bold = False
                run.font.name = 'Times New Roman'
    
    # Traitement des tableaux
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    txt = para.text
                    keys = [k for k in replacements if k in txt]
                    if keys:
                        keys_sorted = sorted(keys, key=len, reverse=True)
                        pattern = '|'.join(map(re.escape, keys_sorted))
                        segs = re.split(f'({pattern})', txt)
                        for run in para.runs: run.text = ""
                        for seg in segs:
                            if seg in replacements:
                                val = str(replacements[seg]) if pd.notnull(replacements[seg]) else ""
                                r = para.add_run(val)
                                r.font.bold = True
                                r.font.name = 'Times New Roman'
                                r.font.size = Pt(9)
                            else:
                                r = para.add_run(seg)
                                r.font.bold = False
                                r.font.name = 'Times New Roman'
                                r.font.size = Pt(9)

def obtenir_points(nicad, df_coords):
    points_df = df_coords[df_coords['nicad'] == nicad].copy()
    if points_df.empty:
        return []
    if 'vertex_index' in points_df.columns:
        points_df = points_df.sort_values('vertex_index')
    
    col_x = 'X' if 'X' in points_df.columns else 'x_centroid'
    col_y = 'Y' if 'Y' in points_df.columns else 'y_centroid'
    
    liste_points = []
    for idx, (_, row) in enumerate(points_df.iterrows(), start=1):
        x_val = row[col_x]
        y_val = row[col_y]
        x_fmt = f"{float(x_val):.2f}" if pd.notnull(x_val) else ""
        y_fmt = f"{float(y_val):.2f}" if pd.notnull(y_val) else ""
        liste_points.append((f"P{idx}", x_fmt, y_fmt))
    return liste_points

def parser_beneficiaires(row):
    prenoms = str(row.get('Prenom', '')).split('\n') if pd.notnull(row.get('Prenom')) else []
    noms = str(row.get('Nom', '')).split('\n') if pd.notnull(row.get('Nom')) else []
    col_piece = 'Numero_piece' if 'Numero_piece' in row.index else 'Num_piece'
    pieces = str(row.get(col_piece, '')).split('\n') if pd.notnull(row.get(col_piece)) else []
    
    max_len = max(len(prenoms), len(noms), len(pieces), 1)
    prenoms = prenoms + [''] * (max_len - len(prenoms))
    noms = noms + [''] * (max_len - len(noms))
    pieces = pieces + [''] * (max_len - len(pieces))
    
    beneficiaires = []
    for i in range(max_len):
        prenom = prenoms[i].strip() if i < len(prenoms) else ''
        nom = noms[i].strip() if i < len(noms) else ''
        cni = pieces[i].strip() if i < len(pieces) else ''
        if prenom or nom:
            beneficiaires.append((prenom, nom, cni))
    return beneficiaires

def set_cell_text(cell, text, font_size=10, bold=False, center=False):
    cell.text = ""
    if cell.paragraphs:
        p = cell.paragraphs[0]
    else:
        p = cell.add_paragraph()
    run = p.add_run(str(text))
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.name = 'Times New Roman'
    if center:
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

def set_table_borders(table):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)
    tblPr.append(tblBorders)
    if tbl.tblPr is None:
        tbl.insert(0, tblPr)

def remplir_tableau_coordonnees(doc, table_index, points):
    import math
    if not points or table_index >= len(doc.tables):
        return
    
    table = doc.tables[table_index]
    while len(table.rows) > 0:
        row = table.rows[-1]
        table._tbl.remove(row._tr)
    
    nb_points = len(points)
    if nb_points <= 15:
        nb_blocs = 1
    elif nb_points <= 30:
        nb_blocs = 2
    else:
        nb_blocs = 3
    
    cols_per_bloc = 3
    total_cols = nb_blocs * cols_per_bloc
    
    while len(table.columns) < total_cols:
        table.add_column(width=Cm(1.5))
    
    header_row = table.add_row()
    for i in range(nb_blocs):
        base_idx = i * cols_per_bloc
        if base_idx < len(header_row.cells):
            set_cell_text(header_row.cells[base_idx], "PT", 8, True, True)
            set_cell_text(header_row.cells[base_idx+1], "X", 8, True, True)
            set_cell_text(header_row.cells[base_idx+2], "Y", 8, True, True)
    
    rows_needed = math.ceil(nb_points / nb_blocs)
    
    for r in range(rows_needed):
        row = table.add_row()
        for b in range(nb_blocs):
            point_idx = r + (b * rows_needed)
            if point_idx < nb_points:
                pt_data = points[point_idx]
                base_col = b * cols_per_bloc
                if base_col + 2 < len(row.cells):
                    set_cell_text(row.cells[base_col], pt_data[0], 7.5, True, True)
                    set_cell_text(row.cells[base_col+1], pt_data[1], 7.5, True, True)
                    set_cell_text(row.cells[base_col+2], pt_data[2], 7.5, True, True)
    
    set_table_borders(table)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

def remplir_tableau_beneficiaires(table, beneficiaires):
    while len(table.rows) > 1:
        row = table.rows[-1]
        table._tbl.remove(row._tr)
    for prenom, nom, cni in beneficiaires:
        row = table.add_row()
        set_cell_text(row.cells[0], prenom, 9, bold=True)
        set_cell_text(row.cells[1], nom, 9, bold=True)
        set_cell_text(row.cells[2], cni, 9, bold=True)
    set_table_borders(table)

# === MAIN ===
def main():
    log("[PYTHON] Démarrage du Python Engine...")

    # Créer les dossiers de sortie
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    output_indiv_dir = os.path.join(OUTPUT_DIR, "Individuelles")
    output_coll_dir = os.path.join(OUTPUT_DIR, "Collectives")
    os.makedirs(output_indiv_dir, exist_ok=True)
    os.makedirs(output_coll_dir, exist_ok=True)

    log("[PYTHON] Chargement Excel...")
    df_indiv = pd.read_excel(FILE_INDIV_DELIB)
    df_indiv['nicad'] = df_indiv['nicad'].apply(clean_id)
    log(f"   ✓ {len(df_indiv)} Délibérations Individuelles")

    df_coll = pd.read_excel(FILE_COLL_DELIB)
    df_coll['nicad'] = df_coll['nicad'].apply(clean_id)
    log(f"   ✓ {len(df_coll)} Délibérations Collectives")

    df_coord_pi = pd.read_excel(FILE_COORD_PI)
    df_coord_pi['nicad'] = df_coord_pi['nicad'].apply(clean_id)
    log(f"   ✓ {len(df_coord_pi)} Coordonnées PI")

    df_coord_pc = pd.read_excel(FILE_COORD_PC)
    df_coord_pc['nicad'] = df_coord_pc['nicad'].apply(clean_id)
    log(f"   ✓ {len(df_coord_pc)} Coordonnées PC")

    # Diagnostic
    nicads_delib = set(df_indiv['nicad'].unique())
    nicads_coords = set(df_coord_pi['nicad'].unique())
    matchs = nicads_delib & nicads_coords
    log(f"   🔍 [DIAGNOSTIC] Correspondance PI: {len(matchs)} / {len(nicads_delib)}")

    # Générer Individuels
    log("\\n📄 Génération Individuelles...")
    nb_gen = 0
    for idx, row in df_indiv.iterrows():
        nicad = row['nicad']
        try:
            doc = Document(TEMPLATE_INDIV)
            nettoyer_connexions_donnees(doc)
            optimiser_mise_en_page(doc)
            
            replacements = {
                '«Prenom»': row.get('Prenom', ''), '«Nom»': row.get('Nom', ''),
                '«nicad»': nicad, '«superficie»': row.get('superficie', ''),
                '«Village»': row.get('Village', ''), '«type_usag»': row.get('type_usag', ''),
                '«Num_piece»': row.get('Num_piece', ''), '«Type_piece»': row.get('Type_piece', ''),
                '«Date_naissance»': row.get('Date_naissance', ''), '«Telephone»': row.get('Telephone', '')
            }
            
            remplacer_texte(doc, replacements)
            optimiser_espacement(doc)
            reduire_texte_legal(doc)
            
            points = obtenir_points(nicad, df_coord_pi)
            if doc.tables and points:
                remplir_tableau_coordonnees(doc, 0, points)
            
            output_path = os.path.join(output_indiv_dir, f"Extrait_PI_{nicad}.docx")
            doc.save(output_path)
            nb_gen += 1
            if nb_gen % 50 == 0:
                log(f"   ... {nb_gen} générés")
        except Exception as e:
            log(f"   ❌ Erreur {nicad}: {str(e)}")
    log(f"   ✓ {nb_gen} Extraits Individuels générés.")

    # Générer Collectifs
    log("\\n📄 Génération Collectives...")
    nb_gen_coll = 0
    for idx, row in df_coll.iterrows():
        nicad = row['nicad']
        try:
            doc = Document(TEMPLATE_COLL)
            nettoyer_connexions_donnees(doc)
            optimiser_mise_en_page(doc)
            
            replacements = {
                '«nicad»': nicad, '«superficie»': row.get('superficie', ''),
                '«Village»': row.get('Village', ''), '«type_usa»': row.get('type_usa', ''),
                '«Num_piece»': row.get('Numero_piece', '')
            }
            
            remplacer_texte(doc, replacements)
            optimiser_espacement(doc)
            reduire_texte_legal(doc)
            
            benefs = parser_beneficiaires(row)
            if len(doc.tables) >= 1:
                remplir_tableau_beneficiaires(doc.tables[0], benefs)
            
            points = obtenir_points(nicad, df_coord_pc)
            if len(doc.tables) >= 2:
                remplir_tableau_coordonnees(doc, 1, points)
            
            output_path = os.path.join(output_coll_dir, f"Extrait_PC_{nicad}.docx")
            doc.save(output_path)
            nb_gen_coll += 1
        except Exception as e:
            log(f"   ❌ Erreur {nicad}: {str(e)}")
    log(f"   ✓ {nb_gen_coll} Extraits Collectifs générés.")

    # ZIP
    log("\\n📦 Création du fichier ZIP...")
    shutil.make_archive("/output/Resultats_Extraits", 'zip', OUTPUT_DIR)
    log("✅ ZIP créé : /output/Resultats_Extraits.zip")

# Appel principal
main()
