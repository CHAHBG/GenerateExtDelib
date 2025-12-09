"""
Script pour créer des modèles de démonstration compatibles avec docxtemplater
Marges étroites (1.27cm) pour tenir sur une page
"""

from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.section import WD_ORIENT

def set_narrow_margins(doc):
    """Set narrow margins (1.27cm = 0.5 inch) for all sections"""
    for section in doc.sections:
        section.top_margin = Cm(1.27)
        section.bottom_margin = Cm(1.27)
        section.left_margin = Cm(1.27)
        section.right_margin = Cm(1.27)

def set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)
    tblPr.append(tblBorders)
    if tbl.tblPr is None:
        tbl.insert(0, tblPr)

def create_template_individuel():
    doc = Document()
    set_narrow_margins(doc)
    
    # Titre (compact)
    title = doc.add_heading('EXTRAIT DE DÉLIBÉRATION', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.size = Pt(14)
    
    # Sous-titre
    sub = doc.add_paragraph("PARCELLE INDIVIDUELLE")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.size = Pt(10)
    sub.runs[0].bold = True
    
    # Infos principales (compact, no extra spacing)
    info_style = {'size': Pt(10), 'space_after': Pt(2)}
    
    def add_field(label, tag):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r1 = p.add_run(f"{label} : ")
        r1.bold = True
        r1.font.size = Pt(10)
        r2 = p.add_run(tag)
        r2.font.size = Pt(10)
    
    add_field("NICAD", "«nicad»")
    add_field("Nom", "«Nom»")
    add_field("Prénom", "«Prenom»")
    add_field("Village", "«Village»")
    add_field("Superficie", "«superficie»")
    add_field("Type d'usage", "«type_usag»")
    add_field("N° Pièce", "«Num_piece»")
    add_field("Type Pièce", "«Type_piece»")
    
    # Section Coordonnées (2 colonnes, compact)
    doc.add_paragraph()
    h2 = doc.add_heading('COORDONNÉES', level=2)
    for run in h2.runs:
        run.font.size = Pt(11)
    
    table = doc.add_table(rows=2, cols=6)
    set_table_borders(table)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Header
    hdr = table.rows[0].cells
    headers = ["PT", "X", "Y", "PT", "X", "Y"]
    for i, cell in enumerate(hdr):
        cell.text = headers[i]
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(8)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Ligne modèle avec boucle split
    row = table.rows[1].cells
    row[0].text = "«#coords_split»«pt1»"
    row[1].text = "«x1»"
    row[2].text = "«y1»"
    row[3].text = "«pt2»"
    row[4].text = "«x2»"
    row[5].text = "«y2»«/coords_split»"
    for cell in row:
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.size = Pt(8)
    
    # Signature
    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph("LE MAIRE")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    doc.save("DEMO_MODELE_INDIVIDUEL.docx")
    print("✓ DEMO_MODELE_INDIVIDUEL.docx créé (marges étroites)")

def create_template_collectif():
    doc = Document()
    set_narrow_margins(doc)
    
    # Titre
    title = doc.add_heading('EXTRAIT DE DÉLIBÉRATION', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.size = Pt(14)
    
    sub = doc.add_paragraph("PARCELLE COLLECTIVE")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.size = Pt(10)
    sub.runs[0].bold = True
    
    def add_field(label, tag):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r1 = p.add_run(f"{label} : ")
        r1.bold = True
        r1.font.size = Pt(10)
        r2 = p.add_run(tag)
        r2.font.size = Pt(10)
    
    add_field("NICAD", "«nicad»")
    add_field("Village", "«Village»")
    add_field("Superficie", "«superficie»")
    add_field("Type d'usage", "«type_usa»")
    
    # Section Bénéficiaires
    doc.add_paragraph()
    h2 = doc.add_heading('BÉNÉFICIAIRES', level=2)
    for run in h2.runs:
        run.font.size = Pt(11)
    
    table_benef = doc.add_table(rows=2, cols=3)
    set_table_borders(table_benef)
    table_benef.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    hdr = table_benef.rows[0].cells
    for i, txt in enumerate(["PRÉNOM", "NOM", "CNI"]):
        hdr[i].text = txt
        hdr[i].paragraphs[0].runs[0].bold = True
        hdr[i].paragraphs[0].runs[0].font.size = Pt(8)
        hdr[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    row = table_benef.rows[1].cells
    row[0].text = "«#beneficiaires»«Prenom»"
    row[1].text = "«Nom»"
    row[2].text = "«CNI»«/beneficiaires»"
    for cell in row:
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.size = Pt(8)
    
    # Section Coordonnées
    doc.add_paragraph()
    h2 = doc.add_heading('COORDONNÉES', level=2)
    for run in h2.runs:
        run.font.size = Pt(11)
    
    table_coords = doc.add_table(rows=2, cols=6)
    set_table_borders(table_coords)
    table_coords.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    hdr = table_coords.rows[0].cells
    for i, txt in enumerate(["PT", "X", "Y", "PT", "X", "Y"]):
        hdr[i].text = txt
        hdr[i].paragraphs[0].runs[0].bold = True
        hdr[i].paragraphs[0].runs[0].font.size = Pt(8)
        hdr[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    row = table_coords.rows[1].cells
    row[0].text = "«#coords_split»«pt1»"
    row[1].text = "«x1»"
    row[2].text = "«y1»"
    row[3].text = "«pt2»"
    row[4].text = "«x2»"
    row[5].text = "«y2»«/coords_split»"
    for cell in row:
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.size = Pt(8)
    
    # Signature
    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph("LE MAIRE")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    doc.save("DEMO_MODELE_COLLECTIF.docx")
    print("✓ DEMO_MODELE_COLLECTIF.docx créé (marges étroites)")

if __name__ == "__main__":
    print("Création des modèles de démonstration (Marges Étroites 1.27cm)...")
    create_template_individuel()
    create_template_collectif()
    print("\n✅ Modèles créés dans le dossier courant !")
